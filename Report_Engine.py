import docx
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt
from docx.shared import RGBColor
import re
import matplotlib.pyplot as plt
import os
import requests
from collections import OrderedDict
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

class AutoReport:

    def current_date(self):
        tDate = datetime.today().strftime('%d/%m/%Y')
        return tDate
        
    def website_header(self, website):
        req = requests.get(website)
        return req.headers

    def header_present(self, head):
        sec = ['Strict-Transport-Security', 'X-Frame-Options', 'X-XSS-Protection', 'X-Content-Type-Options',
                'Content-Security-Policy', 'Referrer-Policy', 'Expect-CT', 'Feature-Policy']
        present = []
    
        for item in head:
            if item in sec:
                present.append(item)
            elif item in [item.lower() for item in sec]:
                present.append(item)

        return present

    def header_missing(self, headPresent):
        sec = ['Strict-Transport-Security', 'X-Frame-Options', 'X-XSS-Protection', 'X-Content-Type-Options',
                'Content-Security-Policy', 'Referrer-Policy', 'Expect-CT', 'Feature-Policy']

        sec = [item.lower() for item in sec]
        headPresent = [item.lower() for item in headPresent]

        return [x for x in sec if x not in headPresent]

    def header_score(self, headP):
        score = ''
        num = len(headP)

        scoreTable = {
            '8':'A+',
            '7':'A+',
            '6':'A',
            '5':'A',
            '4':'B',
            '3':'C',
            '2':'D',
            '1':'F',
            '0':'F'
            }


        for item in scoreTable:
            if str(item) == str(num):
                score = scoreTable[item]
            else:
                pass

        return score

    def risk_assessment(self, head_miss, document):
        identifier = "#PROM_RISKS#"
        
        header_info = {'Strict-Transport-Security':'''HTTP Strict Transport Security (HSTS) is a web security policy mechanism which helps to protect websites against protocol downgrade attacks and cookie hijacking. It allows web servers to declare that web browsers (or other complying user agents) should only interact with it using secure HTTPS connections, and never via the insecure HTTP protocol.
Example: Strict-Transport-Security: max-age=31536000 ; includeSubDomains

References: https://www.owasp.org/index.php/HTTP_Strict_Transport_Security''',
                    'X-Frame-Options':'''The X-Frame-Options response header improves the protection of web applications against clickjacking. It instructs the browser whether the content can be displayed within frames.
Example: X-Frame-Options: deny

References: https://www.owasp.org/index.php/Clickjacking''',
                    'X-XSS-Protection':'''This header enables the cross-site scripting (XSS) filter in your browser. The X-XSS-Protection header has been deprecated by modern browsers and its use can introduce additional security issues on the client side. As such, it is recommended to set the header as X-XSS-Protection: 0 in order to disable the XSS Auditor, and not allow it to take the default behaviour of the browser handling the response.
Example: X-XSS-Protection: 0

References: https://www.owasp.org/index.php/Cross-site_Scripting_(XSS)''',
                    'X-Content-Type-Options':'''Setting this header will prevent the browser from interpreting files as a different MIME type to what is specified in the Content-Type HTTP header (e.g. treating text/plain as text/css).

Example: X-Content-Type-Options: nosniff

References: https://msdn.microsoft.com/en-us/library/gg622941%28v=vs.85%29.aspx''',
                    'Content-Security-Policy':'''A Content Security Policy (CSP) requires careful tuning and precise definition of the policy. If enabled, CSP has significant impact on the way browsers render pages (e.g., inline JavaScript is disabled by default and must be explicitly allowed in the policy). CSP prevents a wide range of attacks, including cross-site scripting and other cross-site injections.

Example: Content-Security-Policy: script-src 'self'

References: https://www.owasp.org/index.php/Content_Security_Policy''',
                    'Referrer-Policy':'''The Referrer-Policy HTTP header governs which referrer information, sent in the Referrer header, should be included with requests made.

Example: Referrer-Policy: no-referrer

References: https://developer.mozilla.org/en-US/docs/Web/HTTP/Headers/Referrer-Policy''',
                    'Expect-CT':'''The Expect-CT header is used by a server to indicate that browsers should evaluate connections to the host for Certificate Transparency compliance.

Example: Expect-CT: max-age=86400, enforce, report-uri="https://foo.example/report"

References: https://tools.ietf.org/html/draft-ietf-httpbis-expect-ct-02''',
                    'Feature-Policy':'''The Feature-Policy header allows developers to selectively enable and disable use of various browser features and APIs.

Example: Feature-Policy: vibrate 'none'; geolocation 'none'

References: https://wicg.github.io/feature-policy/'''}

        table = document.add_table(rows=len(head_miss)+1, cols=2,  style='Table Grid')
        titles = ['HTTP Header','Details']
        check = ''
        string = ''

        cell = table.cell(0, 0)
        cell.text = titles[0]
        cell = table.cell(0, 1)
        cell.text = titles[1]
        
        for i in range(0, len(head_miss)):
            for x in range(0, 2):
                if x == 0:
                    if i == 0:
                        cell = table.cell(i+1, x)
                        for item in header_info:
                            string = str(item).lower()
                            check = str(head_miss[i]).lower()
                            if check == string:
                                cell.text = item
                                string = ''
                                check = ''
                    else:
                        cell = table.cell(i+1, x)
                        for item in header_info:
                            string = str(item)
                            check = str(head_miss[i])
                            if check.lower() == string.lower():
                                cell.text = item
                                string = ''
                                check = ''
                elif x == 1:
                    if i == 0:
                        cell = table.cell(i+1, x)
                        for item in header_info:
                            string = str(item).lower()
                            check = str(head_miss[i]).lower()
                            if check == string:
                                cell.text = header_info[item]
                                string = ''
                                check = ''
                    else:
                        cell = table.cell(i+1, x)
                        for item in header_info:
                            string = str(item)
                            check = str(head_miss[i])
                            if check.lower() == string.lower():
                                cell.text = header_info[item]
                                string = ''
                                check = ''
                else:
                    string = ''
                    check = ''
                    pass

        self.make_rows_bold(table.rows[0])
        self.adjust_cell_text_size(table)
        self.move_table_after(document, table, identifier)
        self.remove_search_phrase(document, identifier)        
    
    def security_assessment(self, identifier, document, website):
        x = self.website_header(website)
        #print(x)
        pres = self.header_present(x)
        #print(pres)
        miss = self.header_missing(pres)
        #print(miss)
        score = self.header_score(pres)
        #print(score)

        self.risk_assessment(miss, document)
    
        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('test', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(24)
        obj_font.name = 'Calibri (Body)'

        if score == "A+" or score == "A":
            obj_font.color.rgb=RGBColor(16,198,74)
        elif score == "B":
            obj_font.color.rgb=RGBColor(17,95,242)
        elif score == "C":
            obj_font.color.rgb=RGBColor(209,179,27)
        elif score == "D" or score == "E" or score == "F":
            obj_font.color.rgb=RGBColor(232,41,19)
        elif score == "R":
            obj_font.color.rgb=RGBColor(147,148,150)
        else:
            pass
    
        #Traversing body text in document and replacing text with dict value if matches dict key

        try:
            for p in document.paragraphs:
                if p.text.find(identifier)>=0:
                    p.text=p.text.replace(identifier,"")
                    p.add_run(score, style='test').bold = True                        
        except AttributeError:
            pass
        
    def adjust_cell_text_size(self, table):
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
                        
    def remove_search_phrase(self, document, search_phrase):
        #Removing individual search_phrase identifier from document
        try:
            for p in document.paragraphs:
                if p.text.find(search_phrase)>=0:
                    p.text=p.text.replace(search_phrase,"")
        except AttributeError:
            pass

    def move_table_after(self, document, table, search_phrase):
        regexp = re.compile(search_phrase)        

        #Moving and inserting table after search_phrase identifier       
        for paragraph in document.paragraphs:
            if paragraph.text and regexp.search(paragraph.text):
                tbl, p = table._tbl, paragraph._p
                p.addnext(tbl)                
                return paragraph                  

    def move_image_after(self, document, image, search_phrase):
        regexp = re.compile(search_phrase)
        for paragraph in document.paragraphs:
            if paragraph.text and regexp.search(paragraph.text):
                #Removing identifier before add_run to insert
                paragraph.text=paragraph.text.replace(search_phrase,"")
                r = paragraph.add_run()
                r.add_picture(image)
                return paragraph

    def make_rows_bold(self, *rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def create_chart(self):
        """
        ACTIVE VULNERABILITIES
        MOST PROMINENT RISKS
        """
        plt.rcParams["font.family"] = "Calibri"
        labels = ['Critical', 'High', 'Medium', 'Low']
        sizes = [self.vuln_types.count(labels[0].upper()),
                 self.vuln_types.count(labels[1].upper()),
                 self.vuln_types.count(labels[2].upper()),
                 self.vuln_types.count(labels[3].upper())]
        
        color_set = ['#AC56AC','#FF0000','#EB6312','#92d050']
        fig1, ax1 = plt.subplots()
        patches, texts, autotexts = ax1.pie(sizes, labels=None, colors=color_set, autopct='%1.1f%%',
                shadow=False, startangle=70, wedgeprops = { 'linewidth' : 1 , 'edgecolor' : 'white'})
        leg = plt.legend(loc='upper left',labels=labels)
        ax1.title.set_text('Active Vulnerabilities')
        ax1.title.set_color("red")

        plt.setp(leg.get_texts(), color='#4d4f4a')

        for text in texts:
            text.remove()
        for autotext in autotexts:
            autotext.set_color("white")
        
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        plt.tight_layout()
        plt.savefig('ActiveVulnerabilities.jpg', dpi=300)

    def _set_cell_background(self, cell, fill, color=None, val=None):
        """
        @fill: Specifies the color to be used for the background
        @color: Specifies the color to be used for any foreground
        pattern specified with the val attribute
        @val: Specifies the pattern to be used to lay the pattern
        color over the background color.
        """


        cell_properties = cell._element.tcPr
        try:
            cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
        except IndexError:
            cell_shading = OxmlElement('w:shd') # add new w:shd element to it
        if fill:
            cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
        if color:
            pass # TODO
        if val:
            pass # TODO
        cell_properties.append(cell_shading)  # finally extend cell props with shading element

    
    def identifiers(self):

        keyText = {"#BUSINESS_NAME#":"Name of business",
                   "#ASSESS_DATE#":"Vulnerability assessment date",
                   "#ASSESS_TYPE#":"Type of Assessment Basic/Standard/Premium",
                   "#EXEC_SUM_COM#":"High-level executive summary of website and vulnerabilities",
                   "#WEBSITE#":"url of website being assessed",
                   "#HIGH_RISK_VULN_COMMENT#":"risk assessment comments of vulnerabilities",
                   "#CLOSING_COM#":"Closing comments and plan to fix vulnerabilities",
                   "#VULN_TITLE_H1#":"High vulnerability title",
                   "#VULN_DETAILS_H1#":"High vulnerability details",
                   "#VULN_LOC_H1#":"High vulnerability location in website or N\A",
                   "#VULN_REM_H1#":"How to remediate\fix High vulnerability",
                   "#VULN_TITLE_M1#":"Medium vulnerability title",
                   "#VULN_DETAILS_M1#":"Medium vulnerability details",
                   "#VULN_LOC_M1#":"Medium vulnerability location in website or N\A",
                   "#VULN_REM_M1#":"How to remediate\fix Medium vulnerability",
                   "#VULN_TITLE_L1#":"Low vulnerability title",
                   "#VULN_DETAILS_L1#":"Low vulnerability details",
                   "#VULN_LOC_L1#":"Low vulnerability location in website or N\A",
                   "#VULN_REM_L1#":"How to remediate\fix Low vulnerability"}

        for item in keyText:
            print(item, ":", keyText[item])
        
    def updateBody(self, document, business, website):

        #Dict of body text identifiers and replacement text
        bodyText = {"#BUSINESS_NAME#":"",
                      "#EXEC_SUM_COM#":"This is the executive summary",
                      "#WEBSITE#":"",
                      "#HIGH_RISK_VULN_COMMENT#":"SQL injection is present and putting website at risk",
                      "#CLOSING_COM#":"Based on these results it is highly recommended to liaise with the creator and host of your website to remediate any vulnerabilities and risks outlined in this report. If you do require any further assistance from Infrasecurity, please contact us."}

        bodyText["#BUSINESS_NAME#"] = business
        bodyText["#WEBSITE#"] = website

        #Traversing body text in document and replacing text with dict value if matches dict key
        for i in bodyText:
            try:
                for p in document.paragraphs:
                    if p.text.find(i)>=0:
                        p.text=p.text.replace(i,bodyText[i])
            except AttributeError:
                pass
                    
    def vulnTable(self, document):
        x = []
        vuln_level = []
        vuln_title = []
        vuln_definition = []
        vuln_fix = []
        identifier = "#VULN_TABLE#"

        #open txt report
        with open("rapidreport.txt", "r") as r:
            for i in r:
                #uncomment to print items in rapidreport.txt list
                #print(i)
                x.append(i.split('%$'))
            r.close()

        #separate list of lists to separate lists
        for i in x:
            vuln_level.append(i[0])
            vuln_title.append(i[1])
            vuln_definition.append(i[2])
            vuln_fix.append(i[3])

        #remove [" from beginning and end of 2 lists
        vuln_level = [x[2:] for x in vuln_level]
        vuln_fix = [x[:-2] for x in vuln_fix]
        vuln_title = [x[1:] for x in vuln_title]
        vuln_definition = [x[1:] for x in vuln_definition]
        vuln_fix = [x[1:] for x in vuln_fix]

        #Insert colomn titles to beginning of all lists
        vuln_level.insert(0, "Importance")
        vuln_title.insert(0, "Vulnerability")
        vuln_definition.insert(0, "Details")
        vuln_fix.insert(0, "Remediation")


        #Create table
        table = document.add_table(rows=len(vuln_level), cols=4,  style='Table Grid')

        #for in range table rows
        for i in range(0, len(vuln_level)):
            #for in range columns
            for x in range(0, 4):
                #Depending on colum add text from specific list
                if x == 0:
                    cell = table.cell(i, x)
                    cell.text = vuln_level[i]
                    #cell.width = 1380744
                    if vuln_level[i] == "HIGH":
                        self._set_cell_background(table.rows[i].cells[x], 'FF0000')
                    elif vuln_level[i] == "MEDIUM":
                        self._set_cell_background(table.rows[i].cells[x], 'EB6312')
                    elif vuln_level[i] == "LOW":
                        self._set_cell_background(table.rows[i].cells[x], 'CFEBAE')
                    elif vuln_level[i] == "INFO":
                        self._set_cell_background(table.rows[i].cells[x], '2EADE8')
                    elif vuln_level[i] == "CRITICAL":
                        self._set_cell_background(table.rows[i].cells[x], 'AC56AC')
                elif x == 1:
                    cell = table.cell(i, x)
                    cell.text = vuln_title[i]
                    #cell.width = 1380744
                elif x == 2:
                    cell = table.cell(i, x)
                    cell.text = vuln_definition[i]
                    #cell.width = 1380744
                elif x == 3:
                    cell = table.cell(i, x)
                    cell.text = vuln_fix[i]
                    #cell.width = 1380744
                else:
                    pass

        self.make_rows_bold(table.rows[0])
        self.adjust_cell_text_size(table)
        self.move_table_after(document, table, identifier)
        self.remove_search_phrase(document, identifier)
        self.vuln_types = vuln_level
        
    def updateTables(self, document):
        assessDate = self.current_date()
        #Dict of table identifiers and replacement text
        tableText = {"#ASSESS_DATE#":assessDate,
                      "#ASSESS_TYPE#":"Standard",
                      "#VULN_TITLE_H1#":"SQL Injection",
                      "#VULN_DETAILS_H1#":"Database vulnerability",
                      "#VULN_LOC_H1#":"N/A",
                      "#VULN_REM_H1#":"Remove identifiers in SQL code",
                      "#VULN_TITLE_M1#":"No HTTPS Certificate found",
                      "#VULN_DETAILS_M1#":"Confidentiality and Integrity issue",
                      "#VULN_LOC_M1#":"N/A",
                      "#VULN_REM_M1#":"Get HTTP Certificate",
                      "#VULN_TITLE_L1#":"Site Map Discovered",
                      "#VULN_DETAILS_L1#":"Full details of website revealed to attacker",
                      "#VULN_LOC_L1#":"N/A",
                      "#VULN_REM_L1#":"Remove site map from database"}

        #Traversing text in cells, cells in tables and tables in documents
        for i in tableText:
            try:
                for t in document.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                #Find and replace text that matches dict key with dict value
                                if p.text.find(i)>=0:
                                    p.text=p.text.replace(i,tableText[i])                        
            except AttributeError:
                pass

       
    def main(self):
        business = 'lockandhandle Ltd'
        website = 'http://www.lockandhandle.com/'
        #website = 'https://www.metoffice.gov.uk/'
        
        #Importing docx file
        document = docx.Document('Infrasecurity Vulnerability Assessment template test 2.docx')
        
        #Setting house style to replacement text
        style = document.styles['Normal']
        font = style.font
        font.size = docx.shared.Pt(11)
        font.name = 'Calibri Light'

        #self.identifiers()
        
        #Search and replace identifiers in body text
        self.updateBody(document, business, website)
        
        #Search and replace identifiers in table cells
        self.updateTables(document)

        #Create vulnerability table from rapidreport
        self.vulnTable(document)

        #Create Vulnerability chart from list in vulnTable function
        self.create_chart()

        #move image after text identifier
        self.move_image_after(document, 'ActiveVulnerabilities.jpg', '#VULN_ASSESS_PIE#')

        #Add security score and associated risk assessment details which generated score
        self.security_assessment('#SECURITY_SCORE#', document, website)

        #save changed document
        document.save('InfrasecAssessmenttest.docx')

        

if __name__ == "__main__":
    ar = AutoReport()
    ar.main()
